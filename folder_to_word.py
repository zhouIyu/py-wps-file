#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件夹图片到Word文档工具
将文件夹中的二级目录作为标题，并将目录内的图片插入到Word文档中
"""

import os
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image
import logging

# 支持的图片格式
SUPPORTED_IMAGE_FORMATS = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'}

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


def is_image_file(file_path):
    """检查文件是否为支持的图片格式"""
    return file_path.suffix.lower() in SUPPORTED_IMAGE_FORMATS


def get_image_size_for_word(image_path, max_width_inches=6):
    """
    获取适合Word文档的图片尺寸
    
    Args:
        image_path (Path): 图片文件路径
        max_width_inches (float): 最大宽度（英寸）
    
    Returns:
        tuple: (width, height) 以英寸为单位
    """
    try:
        with Image.open(image_path) as img:
            width, height = img.size
            
            # 计算长宽比
            aspect_ratio = height / width
            
            # 设置最大宽度
            if width > height:
                # 横向图片
                new_width = min(max_width_inches, 6)
                new_height = new_width * aspect_ratio
            else:
                # 纵向图片
                new_height = min(max_width_inches * aspect_ratio, 8)
                new_width = new_height / aspect_ratio
            
            return new_width, new_height
            
    except Exception as e:
        logger.warning(f"无法获取图片尺寸 {image_path}: {e}")
        return 4, 3  # 默认尺寸


def create_word_document(folder_path, output_path="图片文档.docx"):
    """
    创建Word文档，将文件夹中的二级目录作为标题，图片插入到对应标题下
    
    Args:
        folder_path (str): 源文件夹路径
        output_path (str): 输出Word文档路径
    """
    folder_path = Path(folder_path)
    
    if not folder_path.exists():
        logger.error(f"文件夹不存在: {folder_path}")
        return False
    
    if not folder_path.is_dir():
        logger.error(f"路径不是文件夹: {folder_path}")
        return False
    
    # 创建Word文档
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading(f'{folder_path.name} - 图片文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置标题字号为22号
    for run in title.runs:
        run.font.size = Pt(22)
    
    # 统计信息
    total_directories = 0
    total_images = 0
    
    # 获取所有二级目录
    subdirectories = []
    try:
        for item in folder_path.iterdir():
            if item.is_dir():
                subdirectories.append(item)
    except PermissionError:
        logger.error(f"没有权限访问文件夹: {folder_path}")
        return False
    
    # 按目录名排序
    subdirectories.sort(key=lambda x: x.name)
    
    if not subdirectories:
        logger.warning(f"在 {folder_path} 中没有找到子目录")
        # 添加说明
        doc.add_paragraph("没有找到子目录")
        doc.save(output_path)
        return True
    
    logger.info(f"找到 {len(subdirectories)} 个子目录")
    
    # 先扫描所有目录，收集有图片的目录信息
    valid_directories = []
    for subdir in subdirectories:
        # 获取目录中的所有图片文件
        image_files = []
        try:
            for file_path in subdir.iterdir():
                if file_path.is_file() and is_image_file(file_path):
                    image_files.append(file_path)
        except PermissionError:
            logger.warning(f"没有权限访问目录: {subdir}")
            continue
        
        if image_files:  # 只记录有图片的目录
            valid_directories.append((subdir, len(image_files)))
    
    if not valid_directories:
        logger.warning(f"没有找到包含图片的目录")
        doc.add_paragraph("没有找到包含图片的目录")
        doc.save(output_path)
        return True
    
    # 添加目录页
    toc_heading = doc.add_heading('目录', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置目录标题字号为22号
    for run in toc_heading.runs:
        run.font.size = Pt(22)
    
    # 添加目录项
    for i, (subdir, img_count) in enumerate(valid_directories, 1):
        doc.add_paragraph(f"{i}. {subdir.name}")
    
    # 添加分页符
    doc.add_page_break()
    
    # 遍历每个有效目录
    for subdir, img_count in valid_directories:
        logger.info(f"处理目录: {subdir.name}")
        
        # 重新获取目录中的所有图片文件（排序）
        image_files = []
        try:
            for file_path in subdir.iterdir():
                if file_path.is_file() and is_image_file(file_path):
                    image_files.append(file_path)
        except PermissionError:
            logger.warning(f"没有权限访问目录: {subdir}")
            continue
        
        # 按文件名排序
        image_files.sort(key=lambda x: x.name)
        
        total_directories += 1
        
        # 添加目录标题
        heading = doc.add_heading(subdir.name, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 设置标题字号为22号
        for run in heading.runs:
            run.font.size = Pt(22)
        
        logger.info(f"  找到 {len(image_files)} 张图片")
        
        # 添加图片
        for i, img_path in enumerate(image_files, 1):
            try:
                logger.info(f"    添加图片: {img_path.name}")
                
                # 获取适当的图片尺寸
                width, height = get_image_size_for_word(img_path)
                
                # 插入图片
                paragraph = doc.add_paragraph()
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.add_picture(str(img_path), width=Inches(width), height=Inches(height))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                total_images += 1
                
                # 添加一些间距
                doc.add_paragraph()
                
            except Exception as e:
                logger.error(f"    添加图片失败 {img_path.name}: {e}")
                # 添加错误说明
                doc.add_paragraph(f"无法加载图片: {img_path.name} (错误: {str(e)})")
        
        # 在每个目录后添加分页符（除了最后一个）
        if (subdir, img_count) != valid_directories[-1]:
            doc.add_page_break()
    

    
    # 保存文档
    try:
        doc.save(output_path)
        logger.info(f"Word文档已保存: {Path(output_path).absolute()}")
        logger.info(f"统计信息 - 目录: {total_directories}，图片: {total_images}")
        return True
    except Exception as e:
        logger.error(f"保存Word文档失败: {e}")
        return False


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description="将文件夹中的二级目录作为标题，图片插入到Word文档中")
    parser.add_argument("folder_path", help="源文件夹路径")
    parser.add_argument("-o", "--output", default="图片文档.docx", 
                       help="输出Word文档路径 (默认: 图片文档.docx)")
    
    args = parser.parse_args()
    
    # 检查文件夹是否存在
    if not os.path.exists(args.folder_path):
        print(f"错误: 文件夹不存在: {args.folder_path}")
        sys.exit(1)
    
    # 创建Word文档
    success = create_word_document(args.folder_path, args.output)
    
    if success:
        print(f"\n✓ 成功创建Word文档: {args.output}")
    else:
        print("✗ 创建Word文档失败")
        sys.exit(1)


if __name__ == "__main__":
    main() 